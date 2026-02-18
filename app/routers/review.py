import io
import json
import logging
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse

from app.db import queries as db
from app.middleware.review_auth import (
    ReviewUser,
    create_review_token,
    verify_password,
    verify_review_token,
)
from app.models.review_schemas import (
    EvaluationCreate,
    EvaluationResponse,
    ReviewedArticleCreate,
    ReviewedArticleResponse,
    ReviewLoginRequest,
    ReviewLoginResponse,
    ReviewQueryDetail,
    ReviewQueryListItem,
    ReviewQueryListResponse,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/review", tags=["review"])


# ── Auth ────────────────────────────────────────────────────

@router.post(
    "/auth/login",
    response_model=ReviewLoginResponse,
    summary="Reviewer login",
)
async def review_login(body: ReviewLoginRequest):
    user = await db.get_review_user_by_username(body.username)
    if not user or not user["is_active"]:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    if not verify_password(body.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_review_token(user["id"], user["username"])
    return ReviewLoginResponse(token=token, username=user["username"])


# ── Protected routes ────────────────────────────────────────

@router.get(
    "/queries",
    response_model=ReviewQueryListResponse,
    summary="List queries for review",
)
async def list_queries(
    reviewer: ReviewUser = Depends(verify_review_token),
    page: int = Query(default=1, ge=1),
    per_page: int = Query(default=30, ge=1, le=100),
    review_status: str | None = Query(default=None),
    search: str | None = Query(default=None, max_length=200),
):
    rows, total = await db.list_query_logs_for_review(
        page=page, per_page=per_page,
        review_status=review_status, search=search,
    )
    total_pages = (total + per_page - 1) // per_page
    return ReviewQueryListResponse(
        queries=rows, total=total, page=page,
        per_page=per_page, total_pages=total_pages,
    )


@router.get(
    "/queries/{query_id}",
    response_model=ReviewQueryDetail,
    summary="Get query detail for review",
)
async def get_query_detail(
    query_id: int,
    reviewer: ReviewUser = Depends(verify_review_token),
):
    row = await db.get_query_log_detail(query_id)
    if not row:
        raise HTTPException(status_code=404, detail="Query not found")
    if row.get("review_status") == "excluded":
        raise HTTPException(status_code=403, detail="This query has been excluded from review")

    refs = row.get("references")
    if isinstance(refs, str):
        row["references"] = json.loads(refs)

    evaluation = await db.get_evaluation(query_id)
    latest_article = await db.get_latest_reviewed_article(query_id)

    return ReviewQueryDetail(
        **row,
        evaluation=evaluation,
        latest_article=latest_article,
    )


@router.post(
    "/queries/{query_id}/evaluate",
    response_model=EvaluationResponse,
    summary="Submit evaluation checklist",
)
async def evaluate_query(
    query_id: int,
    body: EvaluationCreate,
    reviewer: ReviewUser = Depends(verify_review_token),
):
    log = await db.get_query_log_detail(query_id)
    if not log:
        raise HTTPException(status_code=404, detail="Query not found")

    checklist_dict = body.checklist.model_dump()
    row = await db.upsert_evaluation(
        query_log_id=query_id,
        reviewer_id=reviewer.id,
        checklist=checklist_dict,
        note=body.note,
    )

    # Auto-set review_status
    all_checked = all(checklist_dict.values())
    status = "approved" if all_checked else "reviewed"
    await db.update_review_status(query_id, status)

    if isinstance(row.get("checklist"), str):
        row["checklist"] = json.loads(row["checklist"])

    return EvaluationResponse(**row)


@router.post(
    "/queries/{query_id}/article",
    response_model=ReviewedArticleResponse,
    summary="Save edited article draft",
)
async def save_article_draft(
    query_id: int,
    body: ReviewedArticleCreate,
    reviewer: ReviewUser = Depends(verify_review_token),
):
    log = await db.get_query_log_detail(query_id)
    if not log:
        raise HTTPException(status_code=404, detail="Query not found")

    row = await db.insert_reviewed_article(
        query_log_id=query_id,
        reviewer_id=reviewer.id,
        title=body.title,
        edited_response=body.edited_response,
    )
    return ReviewedArticleResponse(**row)


@router.get(
    "/queries/{query_id}/article",
    response_model=ReviewedArticleResponse | None,
    summary="Get latest article draft",
)
async def get_article_draft(
    query_id: int,
    reviewer: ReviewUser = Depends(verify_review_token),
):
    article = await db.get_latest_reviewed_article(query_id)
    if not article:
        return None
    return ReviewedArticleResponse(**article)


@router.get(
    "/queries/{query_id}/export/{fmt}",
    summary="Export article as .md or .docx",
)
async def export_article(
    query_id: int,
    fmt: str,
    reviewer: ReviewUser = Depends(verify_review_token),
):
    if fmt not in ("md", "docx"):
        raise HTTPException(status_code=400, detail="Format must be 'md' or 'docx'")

    article = await db.get_latest_reviewed_article(query_id)
    if not article:
        raise HTTPException(status_code=404, detail="No article draft found")

    log = await db.get_query_log_detail(query_id)
    refs = log.get("references", []) if log else []
    if isinstance(refs, str):
        refs = json.loads(refs)

    if fmt == "md":
        return _export_markdown(article, log, refs)
    else:
        return _export_docx(article, log, refs)


def _export_markdown(article: dict, log: dict | None, refs: list) -> StreamingResponse:
    lines = [
        "---",
        f"title: \"{article['title']}\"",
        f"query: \"{log['query_text'] if log else ''}\"",
        f"date: {article['created_at'].isoformat() if hasattr(article['created_at'], 'isoformat') else article['created_at']}",
        f"version: {article['version']}",
        "---",
        "",
        f"# {article['title']}",
        "",
        article["edited_response"],
        "",
    ]

    if refs:
        lines.append("## References")
        lines.append("")
        for ref in refs:
            title = ref.get("title", "Untitled")
            url = ref.get("source_url", "")
            lines.append(f"- [{title}]({url})")
        lines.append("")

    content = "\n".join(lines)
    return StreamingResponse(
        iter([content]),
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="article.md"'},
    )


def _export_docx(article: dict, log: dict | None, refs: list) -> StreamingResponse:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    heading = doc.add_heading(article["title"], level=1)
    for run in heading.runs:
        run.font.size = Pt(18)

    # Metadata
    doc.add_paragraph(
        f"Query: {log['query_text'] if log else 'N/A'}\n"
        f"Version: {article['version']}"
    ).style = doc.styles["Normal"]

    # Body
    for para_text in article["edited_response"].split("\n\n"):
        doc.add_paragraph(para_text.strip())

    # References
    if refs:
        doc.add_heading("References", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Title"
        hdr_cells[1].text = "URL"
        for ref in refs:
            row_cells = table.add_row().cells
            row_cells[0].text = ref.get("title", "Untitled")
            row_cells[1].text = ref.get("source_url", "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="article.docx"'},
    )
